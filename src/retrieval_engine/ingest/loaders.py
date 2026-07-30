"""Document loaders for markdown, plain text, HTML, PDF, and docx.

Three decisions worth stating.

Character offsets are offsets into ``Document.text``, the text this loader produced, not
into the bytes on disk. Front matter is stripped before offsets are computed and line
endings are normalised to ``\\n`` first, because every citation in the system is ultimately
a character range and a half-stripped offset corrupts all of them silently.

Metadata is coerced, not passed through. ``Document.metadata`` accepts only scalars, None,
and lists of strings, while PyYAML happily produces dates and nested mappings from front
matter. Coercing at the boundary means a document with an unquoted date in its front matter
loads instead of failing pydantic validation deep in the pipeline.

``PageSpan`` carries pages for PDFs and headings for the text formats. Both answer the same
question, "where in this document did this chunk come from", which is what a citation needs.
"""

from __future__ import annotations

import hashlib
import re
from datetime import date
from pathlib import Path
from typing import Protocol, runtime_checkable

import yaml

from retrieval_engine.errors import DocumentLoadError, UnsupportedFormatError
from retrieval_engine.models import Document, Metadata, PageSpan

#: A leading YAML block delimited by --- lines.
_FRONT_MATTER = re.compile(r"\A---[ \t]*\r?\n(.*?)\r?\n---[ \t]*(?:\r?\n|\Z)", re.DOTALL)

#: Markdown ATX headings, levels 1 to 3. Deeper headings are body text for our purposes.
_HEADING = re.compile(r"^(#{1,3})[ \t]+(.+?)[ \t]*$", re.MULTILINE)

#: Front-matter key that overrides the derived document id.
_ID_KEYS = ("arxiv_id", "doc_id", "id")

_HTML_BLOCKS = "h1, h2, h3, h4, p, li"
_HTML_HEADINGS = frozenset({"h1", "h2", "h3", "h4"})
_DOCX_HEADING_PREFIX = "Heading"

_PARAGRAPH_SEPARATOR = "\n\n"


@runtime_checkable
class Loader(Protocol):
    """Turns one file into a :class:`Document`."""

    def load(self, path: Path) -> Document:
        """Read ``path`` and return a fully populated document.

        Raises:
            DocumentLoadError: the file is missing, unreadable, or malformed.
        """
        ...


# --------------------------------------------------------------------------------------
# Helpers
# --------------------------------------------------------------------------------------


def _normalise_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _collapse(text: str) -> str:
    return " ".join(text.split())


def _sha256(text: str) -> str:
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def _as_str(value: object) -> str:
    if isinstance(value, date):
        return value.isoformat()
    return str(value)


def _coerce_metadata(raw: object) -> Metadata:
    """Flatten parsed front matter into values :class:`Document` will accept.

    Nested mappings and anything else that does not fit ``MetadataValue`` are dropped
    rather than stringified, because a stringified dict is noise no consumer can use.
    """
    if not isinstance(raw, dict):
        return {}
    result: Metadata = {}
    for key, value in raw.items():
        name = str(key)
        if value is None or isinstance(value, str | int | float | bool):
            result[name] = value
        elif isinstance(value, date):
            result[name] = value.isoformat()
        elif isinstance(value, list | tuple):
            result[name] = [_as_str(item) for item in value]
    return result


def _split_front_matter(text: str) -> tuple[Metadata, str]:
    """Return ``(metadata, body)``, where body has the front matter removed."""
    match = _FRONT_MATTER.match(text)
    if match is None:
        return {}, text
    try:
        parsed = yaml.safe_load(match.group(1))
    except yaml.YAMLError as exc:
        msg = f"invalid YAML front matter: {exc}"
        raise DocumentLoadError(msg) from exc
    return _coerce_metadata(parsed), text[match.end() :]


def _heading_spans(text: str) -> list[PageSpan]:
    """One span per markdown heading, running until the next heading of any level."""
    matches = list(_HEADING.finditer(text))
    spans: list[PageSpan] = []
    for index, match in enumerate(matches):
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        spans.append(
            PageSpan(
                page_number=index + 1,
                start_char=match.start(),
                end_char=end,
                label=_collapse(match.group(2)),
            )
        )
    return spans


def _assemble(parts: list[str], heading_at: dict[int, str]) -> tuple[str, list[PageSpan]]:
    """Join text blocks and turn recorded heading positions into spans.

    ``heading_at`` maps an index in ``parts`` to that heading's label. Offsets are computed
    the same way ``str.join`` lays the parts out, so the spans land exactly on the joined
    text rather than approximately.
    """
    body = _PARAGRAPH_SEPARATOR.join(parts)
    starts: list[int] = []
    cursor = 0
    for part in parts:
        starts.append(cursor)
        cursor += len(part) + len(_PARAGRAPH_SEPARATOR)

    ordered = sorted(heading_at)
    spans: list[PageSpan] = []
    for number, index in enumerate(ordered, start=1):
        next_index = ordered[number] if number < len(ordered) else None
        end = starts[next_index] if next_index is not None else len(body)
        spans.append(
            PageSpan(
                page_number=number,
                start_char=starts[index],
                end_char=end,
                label=heading_at[index],
            )
        )
    return body, spans


def _document_id(path: Path, metadata: Metadata) -> str:
    for key in _ID_KEYS:
        value = metadata.get(key)
        if isinstance(value, str) and value.strip():
            return value.strip()
        if isinstance(value, int | float) and not isinstance(value, bool):
            return str(value)
    return path.stem


def _build(
    path: Path,
    body: str,
    media_type: str,
    metadata: Metadata,
    spans: list[PageSpan],
) -> Document:
    return Document(
        doc_id=_document_id(path, metadata),
        source_path=str(path),
        text=body,
        content_hash=_sha256(body),
        media_type=media_type,
        metadata=metadata,
        page_spans=spans,
    )


def _read_text(path: Path) -> str:
    try:
        raw = path.read_text(encoding="utf-8")
    except FileNotFoundError as exc:
        msg = f"no such file: {path}"
        raise DocumentLoadError(msg) from exc
    except UnicodeDecodeError as exc:
        msg = f"{path} is not valid UTF-8: {exc}"
        raise DocumentLoadError(msg) from exc
    except OSError as exc:
        msg = f"cannot read {path}: {exc}"
        raise DocumentLoadError(msg) from exc
    return _normalise_newlines(raw)


def _require_file(path: Path) -> None:
    if not path.is_file():
        msg = f"no such file: {path}"
        raise DocumentLoadError(msg)


# --------------------------------------------------------------------------------------
# Loaders
# --------------------------------------------------------------------------------------


class MarkdownLoader:
    """Markdown with optional YAML front matter, sectioned by heading."""

    media_type = "text/markdown"

    def load(self, path: Path) -> Document:
        metadata, body = _split_front_matter(_read_text(path))
        return _build(path, body, self.media_type, metadata, _heading_spans(body))


class TextLoader:
    """Plain text. No front matter, no structure to recover."""

    media_type = "text/plain"

    def load(self, path: Path) -> Document:
        return _build(path, _read_text(path), self.media_type, {}, [])


class HtmlLoader:
    """HTML, flattened to block text with headings recorded as spans.

    Scripts, styles, and math are dropped: they are never the answer to a question and
    they pollute both the embedding and any quoted span.
    """

    media_type = "text/html"

    def load(self, path: Path) -> Document:
        from selectolax.lexbor import LexborHTMLParser

        raw = _read_text(path)
        try:
            tree = LexborHTMLParser(raw)
        except Exception as exc:  # selectolax raises builtin exceptions on bad input
            msg = f"cannot parse {path} as HTML: {exc}"
            raise DocumentLoadError(msg) from exc

        for selector in ("script", "style", "math"):
            for node in tree.css(selector):
                node.decompose()

        metadata: Metadata = {}
        title_node = tree.css_first("title")
        if title_node is not None and (title := _collapse(title_node.text() or "")):
            metadata["title"] = title

        parts: list[str] = []
        heading_at: dict[int, str] = {}
        for node in tree.css(_HTML_BLOCKS):
            text = _collapse(node.text() or "")
            if not text:
                continue
            if node.tag in _HTML_HEADINGS:
                heading_at[len(parts)] = text
                if "title" not in metadata:
                    metadata["title"] = text
            parts.append(text)

        body, spans = _assemble(parts, heading_at)
        return _build(path, body, self.media_type, metadata, spans)


class PdfLoader:
    """PDF via pypdf, one :class:`PageSpan` per page."""

    media_type = "application/pdf"

    def load(self, path: Path) -> Document:
        from pypdf import PdfReader
        from pypdf.errors import PyPdfError

        _require_file(path)
        try:
            reader = PdfReader(str(path))
            pages = [_normalise_newlines(page.extract_text() or "") for page in reader.pages]
        except PyPdfError as exc:
            msg = f"cannot parse {path} as PDF: {exc}"
            raise DocumentLoadError(msg) from exc
        except Exception as exc:  # pypdf raises assorted builtins on malformed files
            msg = f"cannot parse {path} as PDF: {type(exc).__name__}: {exc}"
            raise DocumentLoadError(msg) from exc

        metadata: Metadata = {}
        info = reader.metadata
        if info is not None:
            if info.title:
                metadata["title"] = _collapse(str(info.title))
            if info.author:
                metadata["authors"] = [_collapse(str(info.author))]

        parts: list[str] = []
        spans: list[PageSpan] = []
        cursor = 0
        for number, page_text in enumerate(pages, start=1):
            cleaned = page_text.strip()
            if not cleaned:
                continue
            spans.append(
                PageSpan(
                    page_number=number,
                    start_char=cursor,
                    end_char=cursor + len(cleaned),
                    label=f"page {number}",
                )
            )
            parts.append(cleaned)
            cursor += len(cleaned) + len(_PARAGRAPH_SEPARATOR)

        return _build(path, _PARAGRAPH_SEPARATOR.join(parts), self.media_type, metadata, spans)


class DocxLoader:
    """docx via python-docx, sectioned by paragraphs styled as headings."""

    media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    def load(self, path: Path) -> Document:
        import docx
        from docx.opc.exceptions import PackageNotFoundError

        _require_file(path)
        try:
            document = docx.Document(str(path))
        except PackageNotFoundError as exc:
            msg = f"cannot parse {path} as docx: {exc}"
            raise DocumentLoadError(msg) from exc
        except Exception as exc:  # python-docx surfaces zip and xml errors as builtins
            msg = f"cannot parse {path} as docx: {type(exc).__name__}: {exc}"
            raise DocumentLoadError(msg) from exc

        metadata: Metadata = {}
        properties = document.core_properties
        if properties.title:
            metadata["title"] = _collapse(properties.title)
        if properties.author:
            metadata["authors"] = [_collapse(properties.author)]

        parts: list[str] = []
        heading_at: dict[int, str] = {}
        for paragraph in document.paragraphs:
            text = _collapse(paragraph.text)
            if not text:
                continue
            style = paragraph.style
            name = getattr(style, "name", "") or ""
            if name.startswith(_DOCX_HEADING_PREFIX):
                heading_at[len(parts)] = text
                if "title" not in metadata:
                    metadata["title"] = text
            parts.append(text)

        body, spans = _assemble(parts, heading_at)
        return _build(path, body, self.media_type, metadata, spans)


# --------------------------------------------------------------------------------------
# Dispatch
# --------------------------------------------------------------------------------------

_LOADERS: dict[str, Loader] = {
    ".md": MarkdownLoader(),
    ".markdown": MarkdownLoader(),
    ".txt": TextLoader(),
    ".text": TextLoader(),
    ".html": HtmlLoader(),
    ".htm": HtmlLoader(),
    ".pdf": PdfLoader(),
    ".docx": DocxLoader(),
}

#: Extensions the corpus walker will pick up.
SUPPORTED_EXTENSIONS = frozenset(_LOADERS)


def get_loader(path: Path) -> Loader:
    """Return the loader for ``path``.

    Raises:
        UnsupportedFormatError: nothing claims this extension.
    """
    loader = _LOADERS.get(path.suffix.lower())
    if loader is None:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        msg = f"unsupported file type {path.suffix!r} for {path}; supported: {supported}"
        raise UnsupportedFormatError(msg)
    return loader


def load_document(path: Path) -> Document:
    """Load one file, dispatching on its extension."""
    return get_loader(path).load(path)


def iter_source_files(directory: Path) -> list[Path]:
    """Every loadable file under ``directory``, sorted for a deterministic ingest order."""
    if not directory.is_dir():
        msg = f"no such directory: {directory}"
        raise DocumentLoadError(msg)
    return sorted(
        path
        for path in directory.rglob("*")
        if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS
    )


__all__ = [
    "SUPPORTED_EXTENSIONS",
    "DocxLoader",
    "HtmlLoader",
    "Loader",
    "MarkdownLoader",
    "PdfLoader",
    "TextLoader",
    "get_loader",
    "iter_source_files",
    "load_document",
]
