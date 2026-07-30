"""Loaders: real files for every format, and offsets that land where they claim to.

Every fixture here is a real file written to tmp_path, including a byte-level PDF with an
actual text layer, because a loader tested against a mock proves only that the mock works.
"""

from __future__ import annotations

from pathlib import Path

import docx
import pytest

from retrieval_engine.errors import DocumentLoadError, UnsupportedFormatError
from retrieval_engine.ingest.loaders import (
    SUPPORTED_EXTENSIONS,
    DocxLoader,
    HtmlLoader,
    Loader,
    MarkdownLoader,
    PdfLoader,
    TextLoader,
    get_loader,
    iter_source_files,
    load_document,
)

MARKDOWN = """---
arxiv_id: '2401.12345'
title: Hybrid Retrieval
authors:
  - A. Author
  - B. Author
published: 2024-01-15
url: https://arxiv.org/abs/2401.12345
---

# Hybrid Retrieval

Dense and lexical retrieval fail in different ways.

## Fusion

Reciprocal rank fusion combines the two ranked lists without score calibration.

## Reranking

A cross-encoder reorders the shortlist.
"""


def _minimal_pdf(text: str) -> bytes:
    """Build a valid one-page PDF with an extractable text layer.

    Written by hand rather than with a PDF library because none of the runtime
    dependencies can author a text layer, and a PDF fixture whose text cannot be
    extracted would not test the loader at all.
    """
    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R "
        b"/Resources << /Font << /F1 5 0 R >> >> >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    out = bytearray(b"%PDF-1.4\n")
    offsets: list[int] = []
    for number, body in enumerate(objects, start=1):
        offsets.append(len(out))
        out += f"{number} 0 obj\n".encode() + body + b"\nendobj\n"
    xref_at = len(out)
    out += f"xref\n0 {len(objects) + 1}\n".encode()
    out += b"0000000000 65535 f \n"
    for offset in offsets:
        out += f"{offset:010d} 00000 n \n".encode()
    out += (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\nstartxref\n{xref_at}\n%%EOF\n"
    ).encode()
    return bytes(out)


def _write_docx(path: Path) -> None:
    document = docx.Document()
    document.core_properties.title = "Docx Fixture"
    document.core_properties.author = "C. Author"
    document.add_heading("Methods", level=1)
    document.add_paragraph("We measure retrieval quality with nDCG at five.")
    document.add_heading("Results", level=2)
    document.add_paragraph("Hybrid retrieval beat the dense baseline.")
    document.save(str(path))


# --- markdown ---------------------------------------------------------------------------


def test_markdown_front_matter_is_parsed_and_stripped(tmp_path: Path) -> None:
    path = tmp_path / "paper.md"
    path.write_text(MARKDOWN, encoding="utf-8")

    document = load_document(path)

    assert document.metadata["title"] == "Hybrid Retrieval"
    assert document.metadata["authors"] == ["A. Author", "B. Author"]
    assert "arxiv_id" in document.metadata
    # The front matter must not survive into the retrievable text.
    assert "---" not in document.text
    assert document.text.lstrip().startswith("# Hybrid Retrieval")


def test_markdown_doc_id_comes_from_front_matter(tmp_path: Path) -> None:
    path = tmp_path / "whatever-filename.md"
    path.write_text(MARKDOWN, encoding="utf-8")

    assert load_document(path).doc_id == "2401.12345"


def test_doc_id_falls_back_to_the_file_stem(tmp_path: Path) -> None:
    path = tmp_path / "no-front-matter.md"
    path.write_text("# Heading\n\nBody text.\n", encoding="utf-8")

    assert load_document(path).doc_id == "no-front-matter"


def test_unquoted_date_in_front_matter_is_coerced(tmp_path: Path) -> None:
    """PyYAML turns 2024-01-15 into a datetime.date, which the schema would reject."""
    path = tmp_path / "dated.md"
    path.write_text(MARKDOWN, encoding="utf-8")

    document = load_document(path)

    assert document.metadata["published"] == "2024-01-15"
    assert isinstance(document.metadata["published"], str)


def test_nested_mapping_in_front_matter_is_dropped(tmp_path: Path) -> None:
    """A stringified dict is noise no consumer can use, so it is dropped, not coerced."""
    path = tmp_path / "nested.md"
    path.write_text(
        "---\ntitle: Nested\nextra:\n  nested: value\ncount: 3\nflag: true\n---\n\nBody.\n",
        encoding="utf-8",
    )

    document = load_document(path)

    assert "extra" not in document.metadata
    assert document.metadata["title"] == "Nested"
    assert document.metadata["count"] == 3
    assert document.metadata["flag"] is True


def test_invalid_front_matter_yaml_raises(tmp_path: Path) -> None:
    path = tmp_path / "broken.md"
    path.write_text("---\ntitle: [unclosed\n---\n\nBody.\n", encoding="utf-8")

    with pytest.raises(DocumentLoadError, match="invalid YAML front matter"):
        load_document(path)


def test_markdown_heading_spans_land_on_their_headings(tmp_path: Path) -> None:
    """The whole citation story rests on these offsets, so assert them against the text."""
    path = tmp_path / "paper.md"
    path.write_text(MARKDOWN, encoding="utf-8")

    document = load_document(path)

    assert [span.label for span in document.page_spans] == [
        "Hybrid Retrieval",
        "Fusion",
        "Reranking",
    ]
    for span in document.page_spans:
        section = document.text[span.start_char : span.end_char]
        assert section.lstrip().startswith("#")
        assert span.label is not None
        assert span.label in section
    # Spans tile the document from the first heading to the end.
    assert document.page_spans[-1].end_char == len(document.text)


def test_markdown_without_headings_has_no_spans(tmp_path: Path) -> None:
    path = tmp_path / "flat.md"
    path.write_text("Just prose, no headings at all.\n", encoding="utf-8")

    assert load_document(path).page_spans == []


def test_crlf_is_normalised(tmp_path: Path) -> None:
    """Offsets must not shift by one per line depending on the checkout's line endings."""
    path = tmp_path / "crlf.md"
    path.write_bytes(b"---\r\ntitle: CRLF\r\n---\r\n\r\n# Head\r\n\r\nBody.\r\n")

    document = load_document(path)

    assert "\r" not in document.text
    assert document.metadata["title"] == "CRLF"
    assert document.page_spans[0].label == "Head"


# --- plain text -------------------------------------------------------------------------


def test_text_loader_keeps_content_verbatim(tmp_path: Path) -> None:
    path = tmp_path / "notes.txt"
    path.write_text("Line one.\nLine two.\n", encoding="utf-8")

    document = load_document(path)

    assert document.text == "Line one.\nLine two.\n"
    assert document.media_type == "text/plain"
    assert document.metadata == {}


def test_non_utf8_file_raises(tmp_path: Path) -> None:
    path = tmp_path / "latin.txt"
    path.write_bytes(b"\xff\xfe\x00broken")

    with pytest.raises(DocumentLoadError, match="not valid UTF-8"):
        load_document(path)


# --- html -------------------------------------------------------------------------------


def test_html_flattens_blocks_and_records_headings(tmp_path: Path) -> None:
    path = tmp_path / "page.html"
    path.write_text(
        """<html><head><title>Page Title</title>
        <style>.x { color: red }</style></head>
        <body>
          <h1>Introduction</h1>
          <p>Retrieval augmented generation grounds an answer in sources.</p>
          <script>console.log('noise')</script>
          <h2>Method</h2>
          <p>We fuse two ranked lists.</p>
          <ul><li>First point</li></ul>
        </body></html>""",
        encoding="utf-8",
    )

    document = load_document(path)

    assert document.metadata["title"] == "Page Title"
    assert "console.log" not in document.text
    assert "color: red" not in document.text
    assert "First point" in document.text
    assert [span.label for span in document.page_spans] == ["Introduction", "Method"]


def test_html_offsets_land_on_their_headings(tmp_path: Path) -> None:
    path = tmp_path / "page.html"
    path.write_text(
        "<html><body><h1>Alpha</h1><p>Body of alpha.</p>"
        "<h2>Beta</h2><p>Body of beta.</p></body></html>",
        encoding="utf-8",
    )

    document = load_document(path)

    for span in document.page_spans:
        assert span.label is not None
        assert document.text[span.start_char : span.end_char].startswith(span.label)
    assert document.page_spans[-1].end_char == len(document.text)


def test_html_title_falls_back_to_first_heading(tmp_path: Path) -> None:
    path = tmp_path / "untitled.html"
    path.write_text("<html><body><h1>Only Heading</h1><p>Body.</p></body></html>", encoding="utf-8")

    assert load_document(path).metadata["title"] == "Only Heading"


# --- pdf --------------------------------------------------------------------------------


def test_pdf_text_is_extracted_with_a_page_span(tmp_path: Path) -> None:
    path = tmp_path / "paper.pdf"
    path.write_bytes(_minimal_pdf("Hybrid retrieval beats dense only retrieval"))

    document = load_document(path)

    assert "Hybrid retrieval" in document.text
    assert document.media_type == "application/pdf"
    assert len(document.page_spans) == 1
    span = document.page_spans[0]
    assert span.page_number == 1
    assert document.text[span.start_char : span.end_char] == document.text


def test_corrupt_pdf_raises(tmp_path: Path) -> None:
    path = tmp_path / "broken.pdf"
    path.write_bytes(b"this is definitely not a pdf")

    with pytest.raises(DocumentLoadError, match="cannot parse"):
        load_document(path)


def test_missing_pdf_raises(tmp_path: Path) -> None:
    with pytest.raises(DocumentLoadError, match="no such file"):
        load_document(tmp_path / "absent.pdf")


# --- docx -------------------------------------------------------------------------------


def test_docx_headings_and_properties(tmp_path: Path) -> None:
    path = tmp_path / "report.docx"
    _write_docx(path)

    document = load_document(path)

    assert document.metadata["title"] == "Docx Fixture"
    assert document.metadata["authors"] == ["C. Author"]
    assert "nDCG" in document.text
    assert [span.label for span in document.page_spans] == ["Methods", "Results"]
    for span in document.page_spans:
        assert span.label is not None
        assert document.text[span.start_char : span.end_char].startswith(span.label)


def test_corrupt_docx_raises(tmp_path: Path) -> None:
    path = tmp_path / "broken.docx"
    path.write_bytes(b"not a zip archive")

    with pytest.raises(DocumentLoadError, match="cannot parse"):
        load_document(path)


def test_missing_docx_raises(tmp_path: Path) -> None:
    with pytest.raises(DocumentLoadError, match="no such file"):
        load_document(tmp_path / "absent.docx")


# --- dispatch ---------------------------------------------------------------------------


@pytest.mark.parametrize(
    ("suffix", "expected"),
    [
        (".md", MarkdownLoader),
        (".markdown", MarkdownLoader),
        (".txt", TextLoader),
        (".html", HtmlLoader),
        (".htm", HtmlLoader),
        (".pdf", PdfLoader),
        (".docx", DocxLoader),
    ],
)
def test_dispatch_picks_the_right_loader(suffix: str, expected: type) -> None:
    assert isinstance(get_loader(Path(f"x{suffix}")), expected)


def test_extension_matching_is_case_insensitive() -> None:
    assert isinstance(get_loader(Path("PAPER.MD")), MarkdownLoader)


def test_unsupported_extension_raises() -> None:
    with pytest.raises(UnsupportedFormatError, match="unsupported file type"):
        get_loader(Path("data.xyz"))


def test_every_loader_satisfies_the_protocol() -> None:
    for suffix in SUPPORTED_EXTENSIONS:
        assert isinstance(get_loader(Path(f"x{suffix}")), Loader)


def test_missing_file_raises_before_parsing(tmp_path: Path) -> None:
    with pytest.raises(DocumentLoadError, match="no such file"):
        load_document(tmp_path / "gone.md")


# --- corpus walking ---------------------------------------------------------------------


def test_iter_source_files_is_sorted_and_filtered(tmp_path: Path) -> None:
    (tmp_path / "nested").mkdir()
    for name in ("b.md", "a.txt", "nested/c.html", "manifest.json", "notes.rst"):
        (tmp_path / name).write_text("x", encoding="utf-8")

    found = iter_source_files(tmp_path)

    assert [path.name for path in found] == ["a.txt", "b.md", "c.html"]


def test_iter_source_files_on_missing_directory_raises(tmp_path: Path) -> None:
    with pytest.raises(DocumentLoadError, match="no such directory"):
        iter_source_files(tmp_path / "absent")


def test_iter_source_files_on_empty_directory(tmp_path: Path) -> None:
    assert iter_source_files(tmp_path) == []


# --- invariants across formats ----------------------------------------------------------


def test_content_hash_is_stable_across_loads(tmp_path: Path) -> None:
    """Change detection is hash-based, so two loads of one file must agree exactly."""
    path = tmp_path / "paper.md"
    path.write_text(MARKDOWN, encoding="utf-8")

    assert load_document(path).content_hash == load_document(path).content_hash


def test_content_hash_changes_with_content(tmp_path: Path) -> None:
    path = tmp_path / "paper.md"
    path.write_text(MARKDOWN, encoding="utf-8")
    before = load_document(path).content_hash

    path.write_text(MARKDOWN + "\nAn added sentence.\n", encoding="utf-8")

    assert load_document(path).content_hash != before


def test_the_fixture_corpus_loads_cleanly(corpus_dir: Path) -> None:
    """The twelve document fixture corpus is what most other tests are built on."""
    paths = iter_source_files(corpus_dir)

    assert len(paths) == 12
    for path in paths:
        document = load_document(path)
        assert document.text.strip()
        assert len(document.content_hash) == 64
        assert document.metadata["title"]
        assert document.page_spans, f"{path.name} should have heading spans"
        for span in document.page_spans:
            assert 0 <= span.start_char <= span.end_char <= len(document.text)
